from __future__ import annotations
from typing import List, Dict, Any
import os
import re
import tempfile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from sop_generator.utils.document_processor import (
    parse_documents_to_chunks,
    create_enhanced_corpus_summary,
)


def save_uploaded_files(uploaded_files) -> List[str]:
    """Persist Streamlit UploadedFile objects to temporary files and return their paths."""
    if not uploaded_files:
        return []
    tmpdir = tempfile.mkdtemp(prefix="npa_docs_")
    paths: List[str] = []
    for uf in uploaded_files:
        try:
            path = os.path.join(tmpdir, uf.name)
            data = None
            try:
                data = uf.read()
            except Exception:
                try:
                    data = uf.getvalue()
                except Exception:
                    data = None
            if data is None:
                continue
            with open(path, "wb") as f:
                f.write(data)
            paths.append(path)
        except Exception:
            continue
    return paths


def summarize_npa_documents(file_paths: List[str]) -> str:
    """Create a compact technical summary of provided documents for prompting."""
    if not file_paths:
        return ""
    chunks = parse_documents_to_chunks(file_paths)
    return create_enhanced_corpus_summary(chunks)


def extract_structure_outline(file_paths: List[str], max_lines: int = 25) -> str:
    """Extract a likely structure outline from uploaded files by collecting heading-like lines."""
    if not file_paths:
        return ""
    chunks = parse_documents_to_chunks(file_paths)
    text_parts = [ch.get("text", "") for ch in chunks if ch.get("text")]
    text = "\n".join(text_parts)

    lines = [ln.strip() for ln in text.splitlines()]
    candidates: List[str] = []

    heading_patterns = [
        re.compile(r"^\d+[\.)]\s+.+"),
        re.compile(r"^(?:Раздел|Глава|Section)\s+\d+[:\.)-]?\s+.+", re.IGNORECASE),
        re.compile(r"^[A-ZА-Я][A-Za-zА-Яа-я0-9 ,;:\-()]{2,80}$"),
    ]

    for ln in lines:
        if not ln or len(ln) > 160:
            continue
        s = ln.strip()
        if any(p.match(s) for p in heading_patterns):
            candidates.append(s)

    outline: List[str] = []
    seen: set[str] = set()
    for c in candidates:
        key = c.lower()
        if key in seen:
            continue
        seen.add(key)
        outline.append(c)
        if len(outline) >= max_lines:
            break

    if not outline:
        for s in lines:
            if s and len(s) < 80:
                outline.append(s)
                if len(outline) >= max_lines:
                    break

    return "\n".join(outline)


def export_sop_to_docx(meta: Dict[str, str], markdown_text: str) -> str:
    """Export SOP markdown-ish text to a simple DOCX file and return the file path."""
    title = meta.get("title", "СОП")
    number = meta.get("number", "")

    doc = Document()

    head = doc.add_heading(level=0)
    run = head.add_run(f"{title}")
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if number:
        p = doc.add_paragraph(f"Номер: {number}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.lstrip().startswith("- "):
            doc.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
            continue
        stripped = line.lstrip()
        if stripped and stripped[0].isdigit() and ". " in stripped[:6]:
            doc.add_paragraph(stripped, style="List Number")
            continue
        doc.add_paragraph(line)

    out_path = os.path.join(tempfile.gettempdir(), "sop_result.docx")
    doc.save(out_path)
    return out_path 