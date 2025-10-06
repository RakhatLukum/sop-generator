from __future__ import annotations
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt
# Optional reportlab imports for PDF generation
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    PDF_GENERATION_AVAILABLE = True
except ImportError:
    PDF_GENERATION_AVAILABLE = False
    print("Warning: PDF generation disabled due to missing reportlab dependency")
import re
from pathlib import Path
from xml.sax.saxutils import escape


def _is_md_table_separator(line: str) -> bool:
    s = line.strip()
    return s.startswith('|') and set(s.replace('|', '').strip()) <= set('-:—━')


def _clean_md_inline(text: str) -> str:
    # Remove bold/italic markers for DOCX (styling could be enhanced later)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)', r'\1', text)
    return text


def _format_markdown_inline_for_pdf(text: str) -> str:
    formatted = escape(text)
    formatted = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', formatted)
    formatted = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)', r'<i>\1</i>', formatted)
    return formatted.replace('\n', '<br/>')


def _write_markdown_to_docx(document: Document, content: str, sop_title: str = "") -> None:
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        # Skip duplicate top-level title/number embedded in content
        if line.startswith('# '):
            embedded_title = line[2:].strip()
            if sop_title and embedded_title.lower() == sop_title.lower():
                # Skip '# Title' and up to two following blank lines
                i += 1
                skipped = 0
                while i < len(lines) and skipped < 2 and not lines[i].strip():
                    i += 1
                    skipped += 1
                # If the next non-blank line is 'Номер:', skip it too
                if i < len(lines) and lines[i].strip().lower().startswith('номер:'):
                    i += 1
                continue

        # Headings: #, ##, ###, #### -> map to DOCX heading levels
        m = re.match(r'^(#{1,6})\s*(.+)$', line)
        if m:
            hashes = len(m.group(1))
            text = _clean_md_inline(m.group(2).strip())
            # Map: '#' -> level 1 (we already used level 0 for main doc title), '##'->2, etc.
            level = min(max(hashes, 1), 6) - 0  # keep as 1..6
            document.add_heading(text, level=level)
            i += 1
            continue

        # Markdown table block: detect separator row and collect contiguous table lines
        if '|' in line:
            # Collect a block until a blank line, a new header, or a list item
            j = i
            block_lines: List[str] = []
            saw_separator = False
            while j < len(lines):
                next_line = lines[j]
                if not next_line.strip():
                    break
                if re.match(r'^(#{1,6})\s+', next_line) or re.match(r'^\s*([-*•–])\s+.+', next_line):
                    break
                block_lines.append(next_line)
                if _is_md_table_separator(next_line) or ('---' in next_line) or ('━' in next_line):
                    saw_separator = True
                j += 1
            if saw_separator and any('|' in bl for bl in block_lines):
                # Build table data (skip separator-like rows)
                table_data: List[List[str]] = []
                for bl in block_lines:
                    # Skip separator rows such as |---|:---:|---|
                    if _is_md_table_separator(bl) or re.match(r'^\s*\|?\s*[-:━—\s]+\|?(\s*\|[-:━—\s]+\|?)*\s*$', bl):
                        continue
                    if '|' in bl:
                        cells = [c.strip() for c in bl.split('|')]
                        # Trim empty boundary cells
                        if cells and cells[0] == '':
                            cells = cells[1:]
                        if cells and cells[-1] == '':
                            cells = cells[:-1]
                        if cells:
                            table_data.append(cells)
                if table_data:
                    cols = max(len(r) for r in table_data)
                    table = document.add_table(rows=len(table_data), cols=cols)
                    try:
                        table.style = 'Light Grid Accent 1'
                    except Exception:
                        pass
                    for r_idx, row in enumerate(table_data):
                        for c_idx in range(cols):
                            cell_text = _clean_md_inline(row[c_idx] if c_idx < len(row) else '')
                            table.rows[r_idx].cells[c_idx].text = cell_text
                i = j
                continue
            # else fall through to paragraph handling

        # Bulleted lists (-, *, •). Use built-in List Bullet style
        if re.match(r'^\s*([-*•–])\s+.+', line):
            text = re.sub(r'^\s*([-*•–])\s+', '', line).strip()
            p = document.add_paragraph(_clean_md_inline(text))
            try:
                p.style = 'List Bullet'
            except Exception:
                pass
            i += 1
            # Consume subsequent list lines
            while i < len(lines) and re.match(r'^\s*([-*•–])\s+.+', lines[i]):
                sub_text = re.sub(r'^\s*([-*•–])\s+', '', lines[i]).strip()
                sp = document.add_paragraph(_clean_md_inline(sub_text))
                try:
                    sp.style = 'List Bullet'
                except Exception:
                    pass
                i += 1
            continue

        # Regular paragraph: accumulate until blank line
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r'^(#{1,6})\s+', lines[i]):
            # Stop if next block is a table or list indicator to keep paragraphs scoped
            if '|' in lines[i] or re.match(r'^\s*([-*•–])\s+.+', lines[i]):
                break
            para_lines.append(lines[i].rstrip())
            i += 1
        paragraph_text = _clean_md_inline('\n'.join(para_lines))
        document.add_paragraph(paragraph_text)


def _clear_document_content(document: Document) -> None:
    """Remove all existing body, header, and footer content from a Document while preserving styles and page settings."""
    try:
        # Clear body paragraphs
        for p in list(document.paragraphs):
            try:
                p._element.getparent().remove(p._element)
            except Exception:
                pass
        # Clear body tables
        for t in list(document.tables):
            try:
                t._element.getparent().remove(t._element)
            except Exception:
                pass
        # Clear headers and footers content
        for section in list(document.sections):
            try:
                hdr = section.header
                for p in list(hdr.paragraphs):
                    try:
                        p._element.getparent().remove(p._element)
                    except Exception:
                        pass
                for t in list(hdr.tables):
                    try:
                        t._element.getparent().remove(t._element)
                    except Exception:
                        pass
                ftr = section.footer
                for p in list(ftr.paragraphs):
                    try:
                        p._element.getparent().remove(p._element)
                    except Exception:
                        pass
                for t in list(ftr.tables):
                    try:
                        t._element.getparent().remove(t._element)
                    except Exception:
                        pass
            except Exception:
                continue
    except Exception:
        # Best-effort cleanup; ignore if document structure differs
        pass


def populate_docx(document: Document, sop_meta: Dict[str, Any], sections: List[Dict[str, Any]]) -> Document:
    # Ensure the output includes only generated content (no leftover template body)
    _clear_document_content(document)

    title = sop_meta.get("title", "СОП")
    number = sop_meta.get("number", "")

    # Main title
    document.add_heading(f"{title}", level=0)
    if number:
        document.add_paragraph(f"Номер: {number}")

    # If there's a single consolidated section, avoid adding an extra section heading
    if len(sections) == 1:
        body = sections[0].get("content", "")
        _write_markdown_to_docx(document, body, sop_title=title)
        return document

    # Otherwise, render each section with its heading and body
    for idx, section in enumerate(sections, start=1):
        document.add_heading(f"{idx}. {section['title']}", level=1)
        body = section.get("content", "")
        _write_markdown_to_docx(document, body, sop_title=title)
    return document


def export_to_docx(document: Document, output_path: str) -> str:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return str(path)


def export_to_pdf(sections: List[Dict[str, Any]], output_path: str, sop_meta: Dict[str, Any]) -> str:
    if not PDF_GENERATION_AVAILABLE:
        raise ImportError("PDF generation is not available. Please install reportlab: pip install reportlab")
    
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    
    # Register DejaVu fonts for Cyrillic support
    try:
        pdfmetrics.registerFont(TTFont('DejaVu', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
        pdfmetrics.registerFont(TTFont('DejaVu-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
        font_name = 'DejaVu'
        font_bold = 'DejaVu-Bold'
    except:
        # Fallback to default fonts
        font_name = 'Helvetica'
        font_bold = 'Helvetica-Bold'
    
    # Use SimpleDocTemplate for better text handling
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    styles = getSampleStyleSheet()
    
    # Create custom style that handles UTF-8
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=16,
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName=font_bold
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=8,
        spaceBefore=12,
        fontName=font_bold
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6,
        fontName=font_name
    )

    table_header_style = ParagraphStyle(
        'CustomTableHeader',
        parent=body_style,
        fontName=font_bold,
        fontSize=9,
        leading=11,
        spaceBefore=0,
        spaceAfter=0
    )

    table_cell_style = ParagraphStyle(
        'CustomTableCell',
        parent=body_style,
        fontName=font_name,
        fontSize=9,
        leading=11,
        spaceBefore=0,
        spaceAfter=0
    )
    
    story = []
    
    # Add title and number
    title = sop_meta.get("title", "СОП")
    number = sop_meta.get("number", "")
    
    story.append(Paragraph(title, title_style))
    if number:
        story.append(Paragraph(f"Номер: {number}", body_style))
    # Add sections
    for idx, section in enumerate(sections, start=1):
        story.append(Paragraph(f"{idx}. {section['title']}", heading_style))
        content = section.get("content", "").strip()
        if content:
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                para = para.strip('\n')
                if not para.strip():
                    continue

                lines = [line for line in para.split('\n') if line.strip()]
                saw_table_separator = any(
                    _is_md_table_separator(line) or re.match(r'^\s*\|?\s*[-:━—\s]+\|?(\s*\|[-:━—\s]+\|?)*\s*$', line)
                    for line in lines
                )
                has_table_pipes = any('|' in line for line in lines)

                if lines and saw_table_separator and has_table_pipes:
                    table_rows: List[List[str]] = []
                    for line in lines:
                        if _is_md_table_separator(line) or re.match(r'^\s*\|?\s*[-:━—\s]+\|?(\s*\|[-:━—\s]+\|?)*\s*$', line):
                            continue
                        if '|' not in line:
                            table_rows = []
                            break
                        cells = [cell.strip() for cell in line.split('|')]
                        if cells and cells[0] == '':
                            cells = cells[1:]
                        if cells and cells[-1] == '':
                            cells = cells[:-1]
                        if not cells:
                            continue
                        table_rows.append(cells)

                    if table_rows:
                        column_count = max(len(row) for row in table_rows)
                        column_lengths = [1] * column_count
                        for row in table_rows:
                            for col_idx in range(column_count):
                                text = row[col_idx] if col_idx < len(row) else ''
                                column_lengths[col_idx] = max(column_lengths[col_idx], len(text))

                        total_length = sum(column_lengths) or column_count
                        available_width = doc.width
                        col_widths = [available_width * (length / total_length) for length in column_lengths]

                        formatted_rows: List[List[Paragraph]] = []
                        for row_index, row in enumerate(table_rows):
                            formatted_row: List[Paragraph] = []
                            for col_idx in range(column_count):
                                raw_text = row[col_idx] if col_idx < len(row) else ''
                                formatted_text = _format_markdown_inline_for_pdf(raw_text)
                                if not formatted_text.strip():
                                    formatted_text = '&#160;'
                                style = table_header_style if row_index == 0 else table_cell_style
                                formatted_row.append(Paragraph(formatted_text, style))
                            formatted_rows.append(formatted_row)

                        table = Table(formatted_rows, colWidths=col_widths)
                        table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('FONTNAME', (0, 0), (-1, 0), font_bold),
                            ('FONTNAME', (0, 1), (-1, -1), font_name),
                            ('FONTSIZE', (0, 0), (-1, -1), 9),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                            ('TOPPADDING', (0, 0), (-1, -1), 4),
                            ('LEFTPADDING', (0, 0), (-1, -1), 4),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                            ('WORDWRAP', (0, 0), (-1, -1), 'LTR')
                        ]))
                        story.append(table)
                        story.append(Spacer(1, 12))
                        continue

                formatted_para = _format_markdown_inline_for_pdf(para)
                story.append(Paragraph(formatted_para, body_style))
                story.append(Spacer(1, 6))
        story.append(Spacer(1, 12))
    
    # Build PDF
    doc.build(story)
    return str(path) 
