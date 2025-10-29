"""Utilities for exporting SOP content to different formats."""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Dict, List
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


_FONT_CACHE: tuple[str, str] | None = None


def _ensure_pdf_fonts() -> tuple[str, str]:
    """Register TrueType fonts with Cyrillic support and return their names."""
    global _FONT_CACHE
    if _FONT_CACHE:
        return _FONT_CACHE

    registered = set(pdfmetrics.getRegisteredFontNames())

    font_candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/local/share/fonts/DejaVuSans.ttf",
        str(Path(__file__).resolve().parent / "fonts" / "DejaVuSans.ttf"),
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/ARIAL.TTF",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    ]
    bold_candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        "/usr/local/share/fonts/DejaVuSans-Bold.ttf",
        str(Path(__file__).resolve().parent / "fonts" / "DejaVuSans-Bold.ttf"),
        "C:/Windows/Fonts/arialbd.ttf",
        "C:/Windows/Fonts/ARIALBD.TTF",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
    ]

    normal_name = "Helvetica"
    bold_name = "Helvetica-Bold"
    normal_path_used: str | None = None

    def _register(name: str, path: str) -> bool:
        if not path:
            return False
        if name in registered:
            return True
        candidate = Path(path)
        if not candidate.exists():
            return False
        try:
            pdfmetrics.registerFont(TTFont(name, str(candidate)))
            registered.add(name)
            return True
        except Exception:
            return False

    for candidate_path in font_candidates:
        if _register("DejaVuSans", candidate_path):
            normal_name = "DejaVuSans"
            normal_path_used = candidate_path
            break

    for candidate_path in bold_candidates:
        if _register("DejaVuSans-Bold", candidate_path):
            bold_name = "DejaVuSans-Bold"
            break
    else:
        # If bold variant missing but normal is available, rely on the same file.
        if normal_name != "Helvetica" and normal_path_used:
            if _register("DejaVuSans-Bold", normal_path_used):
                bold_name = "DejaVuSans-Bold"

    _FONT_CACHE = (normal_name, bold_name)
    return _FONT_CACHE


def load_template(template_dir: str | os.PathLike[str]) -> tuple[Document, Dict[str, Pt]]:
    """Load a DOCX template if available, otherwise create a blank document."""
    template_dir_path = Path(template_dir)
    template_path = template_dir_path / "base_template.docx"
    if template_path.exists():
        try:
            return Document(str(template_path)), {}
        except Exception:
            pass
    return Document(), {}


def apply_styles(doc: Document, styles: Dict[str, Pt]) -> None:
    """Apply style overrides to the document (no-op placeholder)."""
    _ = (doc, styles)


def populate_docx(doc: Document, meta: Dict[str, str], sections: List[Dict[str, str]]) -> Document:
    title = meta.get("title") or "Стандартная операционная процедура"
    number = meta.get("number")

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if number:
        p = doc.add_paragraph(f"Номер документа: {number}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    equipment = meta.get("equipment")
    if equipment:
        doc.add_paragraph(f"Оборудование/процесс: {equipment}")

    doc.add_page_break()

    for idx, section in enumerate(sections, start=1):
        title_text = section.get("title") or f"Раздел {idx}"
        content = section.get("content", "")
        doc.add_heading(title_text, level=1)
        if content:
            _write_markdown_to_docx(doc, content, sop_title=title)

    return doc


def export_to_docx(doc: Document, destination: str) -> str:
    doc.save(destination)
    return destination


def export_to_pdf(sections: List[Dict[str, str]], destination: str, meta: Dict[str, str]) -> str:
    font_normal, font_bold = _ensure_pdf_fonts()

    path = Path(destination)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(str(path), pagesize=A4)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=16,
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName=font_bold,
    )

    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading1"],
        fontSize=14,
        spaceAfter=8,
        spaceBefore=12,
        fontName=font_bold,
    )

    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=6,
        fontName=font_normal,
    )

    table_header_style = ParagraphStyle(
        "CustomTableHeader",
        parent=body_style,
        fontName=font_bold,
        fontSize=9,
        leading=11,
        spaceBefore=0,
        spaceAfter=0,
    )

    table_cell_style = ParagraphStyle(
        "CustomTableCell",
        parent=body_style,
        fontName=font_normal,
        fontSize=9,
        leading=11,
        spaceBefore=0,
        spaceAfter=0,
    )

    story: List = []

    title = meta.get("title", "СОП")
    number = meta.get("number", "")
    equipment = meta.get("equipment")

    story.append(Paragraph(title, title_style))
    if number:
        story.append(Paragraph(f"Номер: {number}", body_style))
    if equipment:
        story.append(Paragraph(f"Оборудование/процесс: {equipment}", body_style))

    for idx, section in enumerate(sections, start=1):
        story.append(Paragraph(f"{idx}. {section.get('title') or f'Раздел {idx}'}", heading_style))
        content = section.get("content", "").strip()
        if not content:
            continue

        paragraphs = content.split("\n\n")
        for para in paragraphs:
            para = para.strip("\n")
            if not para.strip():
                continue

            lines = [line for line in para.split("\n") if line.strip()]
            saw_table_separator = any(
                _is_md_table_separator(line)
                or re.match(r"^\s*\|?\s*[-:━—\s]+\|?(\s*\|[-:━—\s]+\|?)*\s*$", line)
                for line in lines
            )
            has_table_pipes = any("|" in line for line in lines)

            if lines and saw_table_separator and has_table_pipes:
                table_rows: List[List[str]] = []
                for line in lines:
                    if _is_md_table_separator(line) or re.match(
                        r"^\s*\|?\s*[-:━—\s]+\|?(\s*\|[-:━—\s]+\|?)*\s*$",
                        line,
                    ):
                        continue
                    if "|" not in line:
                        table_rows = []
                        break
                    cells = [cell.strip() for cell in line.split("|")]
                    if cells and cells[0] == "":
                        cells = cells[1:]
                    if cells and cells[-1] == "":
                        cells = cells[:-1]
                    if not cells:
                        continue
                    table_rows.append(cells)

                if table_rows:
                    column_count = max(len(row) for row in table_rows)
                    column_lengths = [1] * column_count
                    for row in table_rows:
                        for col_idx in range(column_count):
                            text = row[col_idx] if col_idx < len(row) else ""
                            column_lengths[col_idx] = max(column_lengths[col_idx], len(text))

                    col_widths, cell_padding = _compute_table_col_widths(column_lengths, doc.width)

                    formatted_rows: List[List[Paragraph]] = []
                    for row_index, row in enumerate(table_rows):
                        formatted_row: List[Paragraph] = []
                        for col_idx in range(column_count):
                            raw_text = row[col_idx] if col_idx < len(row) else ""
                            formatted_text = _format_markdown_inline_for_pdf(raw_text)
                            if not formatted_text.strip():
                                formatted_text = "&#160;"
                            style = table_header_style if row_index == 0 else table_cell_style
                            formatted_row.append(Paragraph(formatted_text, style))
                        formatted_rows.append(formatted_row)

                    table = Table(formatted_rows, colWidths=col_widths, repeatRows=1)
                    table.hAlign = "LEFT"
                    table.setStyle(
                        TableStyle(
                            [
                                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                ("FONTNAME", (0, 0), (-1, 0), font_bold),
                                ("FONTNAME", (0, 1), (-1, -1), font_normal),
                                ("FONTSIZE", (0, 0), (-1, -1), 9),
                                ("BOTTOMPADDING", (0, 0), (-1, 0), max(4.0, cell_padding + 2.0)),
                                ("TOPPADDING", (0, 0), (-1, -1), max(2.0, cell_padding)),
                                ("LEFTPADDING", (0, 0), (-1, -1), cell_padding),
                                ("RIGHTPADDING", (0, 0), (-1, -1), cell_padding),
                                ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                                ("WORDWRAP", (0, 0), (-1, -1), "LTR"),
                            ]
                        )
                    )
                    story.append(table)
                    story.append(Spacer(1, 12))
                    continue

            formatted_para = _format_markdown_inline_for_pdf(para)
            story.append(Paragraph(formatted_para, body_style))
            story.append(Spacer(1, 6))

        story.append(Spacer(1, 12))

    doc.build(story)
    return str(path)


def _is_md_table_separator(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and set(s.replace("|", "").strip()) <= set("-:—━")


def _clean_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"\1", text)
    return text


def _write_markdown_to_docx(document: Document, content: str, sop_title: str = "") -> None:
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        if line.startswith("# "):
            embedded_title = line[2:].strip()
            if sop_title and embedded_title.lower() == sop_title.lower():
                i += 1
                skipped = 0
                while i < len(lines) and skipped < 2 and not lines[i].strip():
                    i += 1
                    skipped += 1
                if i < len(lines) and lines[i].strip().lower().startswith("номер:"):
                    i += 1
                continue

        heading_match = re.match(r"^(#{1,6})\s*(.+)$", line)
        if heading_match:
            hashes = len(heading_match.group(1))
            text = _clean_md_inline(heading_match.group(2).strip())
            level = min(max(hashes, 1), 6)
            document.add_heading(text, level=level)
            i += 1
            continue

        if "|" in line:
            j = i
            block_lines: List[str] = []
            saw_separator = False
            while j < len(lines):
                next_line = lines[j]
                if not next_line.strip():
                    break
                if re.match(r"^(#{1,6})\s+", next_line) or re.match(r"^\s*([-*•–])\s+.+", next_line):
                    break
                block_lines.append(next_line)
                if _is_md_table_separator(next_line) or ("---" in next_line) or ("━" in next_line):
                    saw_separator = True
                j += 1
            if saw_separator and any("|" in bl for bl in block_lines):
                table_data: List[List[str]] = []
                for bl in block_lines:
                    if _is_md_table_separator(bl) or re.match(
                        r"^\s*\|?\s*[-:━—\s]+\|?(\s*\|[-:━—\s]+\|?)*\s*$",
                        bl,
                    ):
                        continue
                    if "|" in bl:
                        cells = [c.strip() for c in bl.split("|")]
                        if cells and cells[0] == "":
                            cells = cells[1:]
                        if cells and cells[-1] == "":
                            cells = cells[:-1]
                        if cells:
                            table_data.append(cells)
                if table_data:
                    cols = max(len(r) for r in table_data)
                    table = document.add_table(rows=len(table_data), cols=cols)
                    try:
                        table.style = "Light Grid Accent 1"
                    except Exception:
                        pass
                    for r_idx, row in enumerate(table_data):
                        for c_idx in range(cols):
                            cell_text = _clean_md_inline(row[c_idx] if c_idx < len(row) else "")
                            table.rows[r_idx].cells[c_idx].text = cell_text
                i = j
                continue

        if re.match(r"^\s*([-*•–])\s+.+", line):
            text = re.sub(r"^\s*([-*•–])\s+", "", line).strip()
            p = document.add_paragraph(_clean_md_inline(text))
            try:
                p.style = "List Bullet"
            except Exception:
                pass
            i += 1
            while i < len(lines) and re.match(r"^\s*([-*•–])\s+.+", lines[i]):
                sub_text = re.sub(r"^\s*([-*•–])\s+", "", lines[i]).strip()
                sp = document.add_paragraph(_clean_md_inline(sub_text))
                try:
                    sp.style = "List Bullet"
                except Exception:
                    pass
                i += 1
            continue

        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,6})\s+", lines[i]):
            if "|" in lines[i] or re.match(r"^\s*([-*•–])\s+.+", lines[i]):
                break
            para_lines.append(lines[i].rstrip())
            i += 1
        paragraph_text = _clean_md_inline("\n".join(para_lines))
        document.add_paragraph(paragraph_text)


def _format_markdown_inline_for_pdf(text: str) -> str:
    formatted = escape(text)
    formatted = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", formatted)
    formatted = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<i>\1</i>", formatted)
    return formatted.replace("\n", "<br/>")


def _compute_table_col_widths(column_lengths: List[int], available_width: float) -> tuple[List[float], float]:
    column_count = len(column_lengths)
    if column_count == 0:
        return [], 0.0

    safe_available = max(float(available_width), 1.0)
    safe_lengths = [max(int(length), 1) for length in column_lengths]

    baseline = min(12.0, safe_available / column_count)
    baseline = max(baseline, 1.0)

    baseline_total = baseline * column_count
    extra_available = max(safe_available - baseline_total, 0.0)

    total_length = sum(safe_lengths)
    if total_length <= 0:
        weights = [1.0 / column_count] * column_count
    else:
        weights = [length / total_length for length in safe_lengths]

    col_widths = [baseline + extra_available * weight for weight in weights]

    difference = safe_available - sum(col_widths)
    if column_count and abs(difference) > 1e-6:
        col_widths[-1] += difference

    min_width = min(col_widths) if col_widths else baseline
    cell_padding = min(4.0, max(0.0, (min_width - 1.0) / 2.0))

    return col_widths, cell_padding
__all__ = [
    "load_template",
    "apply_styles",
    "populate_docx",
    "export_to_docx",
    "export_to_pdf",
]
