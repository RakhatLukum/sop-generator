from __future__ import annotations
from pathlib import Path
from typing import Any, Dict
import json
from docx import Document
from docx.shared import Pt
import os


def load_template(template_dir: str) -> tuple[Document, Dict[str, Any]]:
    base = Path(template_dir)
    docx_path = base / "sop_template.docx"
    styles_path = base / "styles.json"

    doc: Document | None = None

    # Preferred template in templates directory
    if docx_path.exists():
        try:
            doc = Document(str(docx_path))
        except Exception:
            doc = None

    # Fallback: allow environment override
    if doc is None:
        env_path = (Path(str(Path.cwd())) / (os.getenv("SOP_DOCX_TEMPLATE_PATH") or "")).resolve() if os.getenv("SOP_DOCX_TEMPLATE_PATH") else None  # type: ignore
        if env_path and env_path.exists() and env_path.suffix.lower() == ".docx":
            try:
                doc = Document(str(env_path))
            except Exception:
                doc = None

    # Fallback: search project root for any plausible DOCX template (e.g., provided sample)
    if doc is None:
        try:
            project_root = Path(__file__).resolve().parents[2]
            candidates: list[Path] = []
            for p in project_root.glob("*.docx"):
                # Prefer filenames that look like structure templates
                if any(key in p.name.lower() for key in ["структ", "template", "sop"]):
                    candidates.append(p)
            if not candidates:
                # As last resort, any docx at root
                candidates = list(project_root.glob("*.docx"))
            if candidates:
                doc = Document(str(sorted(candidates)[0]))
        except Exception:
            doc = None

    if doc is None:
        doc = Document()

    styles: Dict[str, Any] = {}
    if styles_path.exists():
        try:
            styles = json.loads(styles_path.read_text(encoding="utf-8"))
        except Exception:
            styles = {}
    return doc, styles


def apply_styles(document: Document, styles: Dict[str, Any]) -> None:
    if not styles:
        return
    # Basic example: set default font
    default = styles.get("default_font")
    if default:
        for style in document.styles:
            try:
                if hasattr(style, "font"):
                    if default.get("name"):
                        style.font.name = default["name"]
                    if default.get("size"):
                        size_val = default["size"]
                        # Wrap numeric values using Pt for python-docx
                        if isinstance(size_val, (int, float)):
                            style.font.size = Pt(size_val)
                        else:
                            style.font.size = size_val
            except Exception:
                continue 